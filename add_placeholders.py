import docx

doc = docx.Document('HotelEase_SRS.docx')

# 1. Multi-step booking wizard placeholder
in_booking = False
for p in doc.paragraphs:
    if "Transaction Name: Online Room Booking" in p.text:
        in_booking = True
    elif in_booking and "Wireframe" in p.text:
        p.insert_paragraph_before("(To be filled — Insert Activity Diagram for Multi-Step Booking Wizard)", style='Normal')
        in_booking = False

# 2. Payment proof flow placeholder
in_payment = False
for p in doc.paragraphs:
    if "Transaction Name: Payment Processing" in p.text:
        in_payment = True
    elif in_payment and "Wireframe" in p.text:
        p.insert_paragraph_before("(To be filled — Insert Activity Diagram for Payment Proof Flow)", style='Normal')
        in_payment = False

# 3. FO-mediated cancellation transaction
for p in doc.paragraphs:
    if "Transaction Name: Guest Check-In and Check-Out" in p.text:
        p.insert_paragraph_before("Transaction Name: FO-Mediated Cancellation", style='Heading 3')
        p.insert_paragraph_before("Use Case Diagram", style='List Bullet')
        p.insert_paragraph_before("(To be filled — Insert Use Case Diagram for FO-Mediated Cancellation)", style='Normal')
        p.insert_paragraph_before("Use Case Description", style='List Bullet')
        p.insert_paragraph_before("Guest requests cancellation of an Approved booking. FO reviews and approves or rejects it.", style='Body Text')
        p.insert_paragraph_before("Activity Diagram", style='List Bullet')
        p.insert_paragraph_before("(To be filled — Insert Activity Diagram for FO-Mediated Cancellation Flow)", style='Normal')
        p.insert_paragraph_before("Wireframe", style='List Bullet')
        break

# 4. Training mode session flow placeholder
in_training = False
for p in doc.paragraphs:
    if "Transaction Name: Training Mode Management" in p.text:
        in_training = True
    elif in_training and "(To be filled — Insert Activity Diagram for Training Mode)" in p.text:
        p.text = "(To be filled — Insert Activity Diagram for Training Mode Session Flow)"
        in_training = False

doc.save('HotelEase_SRS.docx')
print("Successfully added placeholders")
