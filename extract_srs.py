import docx

doc = docx.Document('HotelEase_SRS.docx')
with open('srs_extracted.txt', 'w', encoding='utf-8') as f:
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip():
            f.write(f"[{i}] {p.style.name}: {p.text}\n")

    f.write("\n\n--- TABLES ---\n\n")
    for t_idx, table in enumerate(doc.tables):
        f.write(f"Table {t_idx}:\n")
        for r_idx, row in enumerate(table.rows):
            row_data = [cell.text.replace('\n', ' ').strip() for cell in row.cells]
            f.write(f"  Row {r_idx}: | " + " | ".join(row_data) + " |\n")
        f.write("\n")
