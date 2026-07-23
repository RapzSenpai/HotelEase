import docx

doc = docx.Document('HotelEase_SRS.docx')

# 1. Tech Stack (Software Interfaces) Updates
# Add Framer Motion, @dnd-kit, React Hook Form, Sonner, Geist font to Frontend Stack
for i, p in enumerate(doc.paragraphs):
    if "React 19 — Component-based UI library" in p.text:
        # We will just append the new libraries to the end of the frontend stack list
        pass
    if "Backend / Cloud Services:" in p.text:
        # Insert before this paragraph
        p.insert_paragraph_before("Framer Motion — Animation library for UI transitions and micro-interactions.", style='List Bullet')
        p.insert_paragraph_before("@dnd-kit/core + utilities — Drag-and-drop for housekeeping kanban.", style='List Bullet')
        p.insert_paragraph_before("React Hook Form — Form state management.", style='List Bullet')
        p.insert_paragraph_before("Sonner — Toast notifications.", style='List Bullet')
        p.insert_paragraph_before("Geist Variable — Typography font.", style='List Bullet')

# 2. User Roles Updates
for i, p in enumerate(doc.paragraphs):
    if p.text.strip() == "Cancel pending or approved bookings.":
        p.text = "Self-cancel a Pending booking (no FO involvement) or request cancellation of an Approved booking (FO-mediated; max 3 lifetime cancellations enforced)."
    if p.text.strip() == "Record payments (GCash, Cash, Check, Credit Card) with folio tracking.":
        p.text = "Record payments (GCash, Bank Transfer, Credit/Debit Card, Over-the-Counter) with folio tracking, balance overpayment validation, and payment proof verification."

# 3. Core Modules Updates
for i, p in enumerate(doc.paragraphs):
    if "Booking Lifecycle: Full workflow from Pending" in p.text:
        p.text = "Booking Lifecycle: Full workflow from Pending → Approved → Checked In → Checked Out → Cancelled, with Front Office (FO) staff managing approvals, check-ins, check-outs, and a new FO-mediated cancellation flow (Cancellation Requested state)."
    if "Payment Processing: Support for GCash, Cash, Check, and Credit Card payments" in p.text:
        p.text = "Payment Processing: Support for 4 methods (GCash, Bank Transfer, Credit/Debit Card, Over-the-Counter). Features proof-required vs proof-exempt logic, Full/Partial payment types, 48-hr payment deadline with auto-expiry, balance overpayment validation, folio management, and PDF receipt generation via jsPDF."

# 4. Security Measures (Password Validation)
for i, p in enumerate(doc.paragraphs):
    if p.text.strip() == "Date validation prevents check-out before check-in and detects overlapping bookings.":
        # Insert password validation right after this bullet
        p_next = p.insert_paragraph_before("Password validation: Enforced in RegisterPage (min 8 characters, at least 1 number).", style='List Bullet')
        # Since insert_paragraph_before inserts before, we actually insert it before the next paragraph "Data Backup and Recovery:"
    if p.text.strip() == "Data Backup and Recovery:":
        p.insert_paragraph_before("Password validation: Enforced in RegisterPage (min 8 characters, at least 1 number).", style='List Bullet')

# 5. Known Limitations
doc.add_heading('Known Limitations', level=2)
limitations = [
    "No real payment gateway. Uses manual proof-upload + FO verification model.",
    "Groq API key exposed client-side (dangerouslyAllowBrowser: true).",
    "No Cloud Functions (Firebase Spark plan limitation). Auto-expiry uses lazy client-side checking.",
    "Testimonials, messages, and announcements are NOT training-mode sandboxed.",
    "Partial real-time sync. Several FO operational pages use one-shot data fetch instead of onSnapshot.",
    "No SMS notifications. In-app and email only.",
    "No multi-language or multi-currency support.",
    "No automated tests or CI pipeline.",
    "No real-time conflict prevention during booking wizard (checked at creation time, race condition possible)."
]
for lim in limitations:
    doc.add_paragraph(lim, style='List Bullet')

# 6. Future Enhancements
doc.add_heading('Future Enhancements', level=2)
enhancements = [
    "Real payment gateway integration (PayMongo, GCash API, Stripe).",
    "Serverless email and background job migration (Cloud Functions).",
    "Real-time sync for remaining one-shot FO pages.",
    "Automated refund and cancellation-window logic.",
    "Groq API key proxy (Vercel Edge Function / Cloudflare Worker).",
    "Training-mode sandboxing for testimonials, messages, and announcements.",
    "Multi-language support.",
    "Composite Firestore index optimizations.",
    "Mobile-native application (React Native or Flutter)."
]
for enh in enhancements:
    doc.add_paragraph(enh, style='List Bullet')

doc.save('HotelEase_SRS.docx')
print("Successfully updated HotelEase_SRS.docx")
